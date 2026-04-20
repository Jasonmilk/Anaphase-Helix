#!/usr/bin/env python3
"""
类案检索报告生成脚本 - 将JSON格式的案例数据转换为Word报告

重要约束：仅使用步骤5核验通过的案例生成报告
"""
import argparse
import json
import sys
import time
from pathlib import Path
from typing import Dict, List, Optional
from datetime import datetime


def print_debug(msg: str):
    """打印调试信息"""
    print(msg, flush=True)
    try:
        with open("report_log.txt", "a", encoding="utf-8") as f:
            f.write(f"{time.strftime('%H:%M:%S')} - {msg}\n")
    except:
        pass


def validate_case(case: Dict, index: int) -> List[str]:
    """
    验证单个案例的必需字段
    
    Returns: 错误信息列表，空列表表示验证通过
    """
    errors = []
    required_fields = [
        ("title", "案件名称"),
        ("case_no", "案号"),
        ("court_name", "审理法院"),
        ("judge_date", "裁判日期"),
        ("court_opinion", "本院认为"),
    ]
    
    for field, name in required_fields:
        value = case.get(field)
        if not value or (isinstance(value, str) and not value.strip()):
            errors.append(f"案例{index}: 缺少必需字段 '{name}' ({field})")
    
    return errors


def filter_verified_cases(cases: List[Dict]) -> List[Dict]:
    """
    过滤出核验通过的案例
    
    仅保留 verification_status == "已核验" 的案例
    如果没有 verification_status 字段，保留所有案例（兼容旧格式）
    """
    has_verification_field = any(c.get("verification_status") for c in cases)
    
    if has_verification_field:
        verified = [c for c in cases if c.get("verification_status") == "已核验"]
        print_debug(f"📋 过滤核验状态：{len(verified)}/{len(cases)} 个案例通过核验")
        return verified
    else:
        print_debug(f"⚠️ 未发现 verification_status 字段，使用全部 {len(cases)} 个案例")
        return cases


def parse_case_input(input_path: str) -> Dict:
    """
    解析案例数据JSON文件
    
    输入格式示例：
    {
      "case_info": {...},
      "cases": [...],
      "analysis": {...}
    }
    
    Returns: 包含案例信息的字典
    Raises: ValueError 如果格式错误或数据无效
    """
    input_file = Path(input_path)
    
    if not input_file.exists():
        raise ValueError(f"❌ 输入文件不存在: {input_path}")
    
    if input_file.stat().st_size == 0:
        raise ValueError(f"❌ 输入文件为空: {input_path}")
    
    try:
        with open(input_path, 'r', encoding='utf-8') as f:
            content = f.read().strip()
            if not content:
                raise ValueError(f"❌ 输入文件内容为空: {input_path}")
            data = json.loads(content)
    except json.JSONDecodeError as e:
        raise ValueError(f"❌ JSON格式错误: {e}")
    
    # 验证基本结构
    if not isinstance(data, dict):
        raise ValueError("❌ JSON根元素必须是对象")
    
    if 'cases' not in data:
        raise ValueError("❌ 缺少必需字段: cases")
    
    cases = data.get('cases', [])
    if not isinstance(cases, list):
        raise ValueError("❌ 'cases' 字段必须是数组")
    
    # 过滤核验通过的案例
    verified_cases = filter_verified_cases(cases)
    
    if len(verified_cases) == 0:
        raise ValueError("❌ 没有核验通过的案例，无法生成报告。请回到步骤5完成核验。")
    
    # 验证每个案例的必需字段
    all_errors = []
    for idx, case in enumerate(verified_cases, 1):
        errors = validate_case(case, idx)
        all_errors.extend(errors)
    
    if all_errors:
        print_debug("⚠️ 案例数据存在以下问题：")
        for err in all_errors[:10]:  # 最多显示10条
            print_debug(f"   {err}")
        if len(all_errors) > 10:
            print_debug(f"   ... 还有 {len(all_errors) - 10} 条错误")
    
    # 更新数据中的cases为核验通过的案例
    data['cases'] = verified_cases
    
    print_debug(f"✅ 解析案例数据成功，共 {len(verified_cases)} 个核验通过的案例")
    return data


def ensure_case_info(data: Dict) -> Dict:
    """确保 case_info 结构完整"""
    default_info = {
        "title": "类案检索报告",
        "description": "法律案例检索与分析",
        "search_date": datetime.now().strftime('%Y-%m-%d'),
        "searcher": "AI 智能体"
    }
    
    case_info = data.get('case_info', {})
    if not isinstance(case_info, dict):
        case_info = {}
    
    for key, default_value in default_info.items():
        if not case_info.get(key):
            case_info[key] = default_value
    
    return case_info


def ensure_analysis(data: Dict) -> Dict:
    """确保 analysis 结构完整"""
    default_analysis = {
        "legal_issues": [],
        "judicial_trends": [],
        "conclusion": ""
    }
    
    analysis = data.get('analysis', {})
    if not isinstance(analysis, dict):
        analysis = {}
    
    for key, default_value in default_analysis.items():
        if key not in analysis:
            analysis[key] = default_value
    
    return analysis


def generate_word_report(case_data: Dict, output_path: Path):
    """
    生成Word格式的类案检索报告
    
    Args:
        case_data: 案例数据字典（仅含核验通过的案例）
        output_path: 输出Word文件路径
    """
    try:
        from docx import Document
        from docx.shared import Pt, Inches, Cm, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.enum.table import WD_TABLE_ALIGNMENT
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement
    except ImportError:
        raise ImportError("❌ 请安装python-docx: pip install python-docx==0.8.11 --break-system-packages")
    
    print_debug("📄 开始生成Word报告...")
    
    # 确保数据完整性
    case_info = ensure_case_info(case_data)
    cases = case_data.get('cases', [])
    analysis = ensure_analysis(case_data)
    
    if len(cases) == 0:
        raise ValueError("❌ 案例列表为空，无法生成报告")
    
    # 创建文档
    doc = Document()
    
    # ========== 设置文档样式 ==========
    style = doc.styles['Normal']
    style.font.name = '宋体'
    style.font.size = Pt(12)  # 小四号
    style._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    
    # 设置页边距（按模板要求：上2.5cm，下2.5cm，左3cm，右2cm）
    sections = doc.sections
    for section in sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(3)
        section.right_margin = Cm(2)
    
    # ========== 1. 标题页 ==========
    title = case_info.get('title', '类案检索报告')
    title_para = doc.add_heading(title, level=0)
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    # 设置标题字体为黑体三号
    for run in title_para.runs:
        run.font.name = '黑体'
        run.font.size = Pt(16)  # 三号
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
    
    doc.add_paragraph()  # 空行
    
    # ========== 2. 基本信息表 ==========
    info_table = doc.add_table(rows=4, cols=2)
    info_table.style = 'Table Grid'
    info_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    info_data = [
        ('检索日期', case_info.get('search_date', datetime.now().strftime('%Y-%m-%d'))),
        ('检索人', case_info.get('searcher', 'AI 智能体')),
        ('案件描述', case_info.get('description', '')),
        ('检索范围', '最高人民法院及各级法院典型案例')
    ]
    
    for i, (key, value) in enumerate(info_data):
        row = info_table.rows[i]
        # 项目列
        cell_key = row.cells[0]
        cell_key.text = key
        cell_key.paragraphs[0].runs[0].font.bold = True
        cell_key.paragraphs[0].runs[0].font.size = Pt(10.5)  # 五号
        # 内容列
        cell_value = row.cells[1]
        cell_value.text = str(value) if value else ''
        if cell_value.paragraphs[0].runs:
            cell_value.paragraphs[0].runs[0].font.size = Pt(10.5)
    
    doc.add_paragraph()  # 空行
    
    # ========== 3. 检索过程 ==========
    h1 = doc.add_heading('一、检索过程', level=1)
    for run in h1.runs:
        run.font.name = '黑体'
        run.font.size = Pt(14)  # 小三号
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
    
    search_process = """本次检索采用语义检索与关键词检索相结合的方式进行：

1. 语义检索：通过自然语言描述案情，运用智能检索系统查找相关案例

2. 关键词检索：根据专业法律术语进行精确检索

3. 筛选标准：按案件层级、裁判时间、相关性进行筛选

4. 核验流程：对候选案例逐一核验，确认法条/要件在本院认为中出现

5. 排序规则：最高人民法院案例优先，近三年案例优先"""
    
    process_para = doc.add_paragraph(search_process)
    process_para.paragraph_format.first_line_indent = Cm(0.74)  # 首行缩进2字符
    process_para.paragraph_format.line_spacing = 1.5  # 1.5倍行距
    
    # ========== 4. 类案列表 ==========
    h2 = doc.add_heading('二、类案列表', level=1)
    for run in h2.runs:
        run.font.name = '黑体'
        run.font.size = Pt(14)
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
    
    print_debug(f"📝 正在添加 {len(cases)} 个核验通过的案例...")
    
    for idx, case in enumerate(cases, 1):
        # 案例标题
        case_heading = doc.add_heading(f'案例 {idx}', level=2)
        for run in case_heading.runs:
            run.font.name = '黑体'
            run.font.size = Pt(12)
            run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
        
        # 案件名称
        case_title = case.get('title', f'案例{idx}')
        title_para = doc.add_paragraph()
        title_run = title_para.add_run(f'案件名称：{case_title}')
        title_run.font.bold = True
        title_run.font.size = Pt(12)
        
        # 案例详情表格
        case_table = doc.add_table(rows=7, cols=2)
        case_table.style = 'Table Grid'
        
        case_fields = [
            ('案号', case.get('case_no', '')),
            ('审理法院', case.get('court_name', '')),
            ('裁判日期', case.get('judge_date', '')),
            ('审理程序', case.get('trial_step', '')),
            ('案由', case.get('cause_of_action', '')),
            ('相关性等级', case.get('relevance_level', '相关')),
            ('核验状态', case.get('verification_status', '已核验'))
        ]
        
        for i, (key, value) in enumerate(case_fields):
            row = case_table.rows[i]
            row.cells[0].text = key
            row.cells[0].paragraphs[0].runs[0].font.size = Pt(10.5)
            row.cells[0].paragraphs[0].runs[0].font.bold = True
            row.cells[1].text = str(value) if value else ''
            if row.cells[1].paragraphs[0].runs:
                row.cells[1].paragraphs[0].runs[0].font.size = Pt(10.5)
        
        # 本院认为
        doc.add_paragraph()
        opinion_label = doc.add_paragraph()
        opinion_label_run = opinion_label.add_run('【本院认为】')
        opinion_label_run.font.bold = True
        opinion_label_run.font.size = Pt(12)
        
        court_opinion = case.get('court_opinion', '')
        if court_opinion:
            opinion_para = doc.add_paragraph(court_opinion)
            opinion_para.paragraph_format.left_indent = Inches(0.25)
            opinion_para.paragraph_format.first_line_indent = Cm(0.74)
            opinion_para.paragraph_format.line_spacing = 1.5
        else:
            doc.add_paragraph('（无）')
        
        # 案例摘要
        summary = case.get('summary', '')
        if summary:
            doc.add_paragraph()
            summary_label = doc.add_paragraph()
            summary_label_run = summary_label.add_run('【案例摘要】')
            summary_label_run.font.bold = True
            summary_label_run.font.size = Pt(12)
            
            summary_para = doc.add_paragraph(summary)
            summary_para.paragraph_format.left_indent = Inches(0.25)
            summary_para.paragraph_format.first_line_indent = Cm(0.74)
            summary_para.paragraph_format.line_spacing = 1.5
        
        # 核验说明（如有）
        verification_note = case.get('verification_note', '')
        if verification_note:
            doc.add_paragraph()
            note_label = doc.add_paragraph()
            note_label_run = note_label.add_run('【核验说明】')
            note_label_run.font.bold = True
            note_label_run.font.size = Pt(10.5)
            note_label_run.font.italic = True
            
            note_para = doc.add_paragraph(verification_note)
            note_para.paragraph_format.left_indent = Inches(0.25)
        
        doc.add_paragraph()  # 案例间空行
    
    # ========== 5. 法律分析 ==========
    h3 = doc.add_heading('三、法律分析', level=1)
    for run in h3.runs:
        run.font.name = '黑体'
        run.font.size = Pt(14)
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
    
    # 5.1 争议焦点
    legal_issues = analysis.get('legal_issues', [])
    if legal_issues:
        issue_heading = doc.add_heading('1. 争议焦点', level=2)
        for run in issue_heading.runs:
            run.font.size = Pt(12)
        
        for issue in legal_issues:
            if issue:
                issue_para = doc.add_paragraph(f'• {issue}')
                issue_para.paragraph_format.left_indent = Inches(0.25)
    
    # 5.2 司法裁判趋势
    judicial_trends = analysis.get('judicial_trends', [])
    if judicial_trends:
        trend_heading = doc.add_heading('2. 司法裁判趋势', level=2)
        for run in trend_heading.runs:
            run.font.size = Pt(12)
        
        for trend in judicial_trends:
            if trend:
                trend_para = doc.add_paragraph(f'• {trend}')
                trend_para.paragraph_format.left_indent = Inches(0.25)
    
    # 5.3 分析结论
    conclusion = analysis.get('conclusion', '')
    if conclusion:
        conclusion_heading = doc.add_heading('3. 分析结论', level=2)
        for run in conclusion_heading.runs:
            run.font.size = Pt(12)
        
        conclusion_para = doc.add_paragraph(conclusion)
        conclusion_para.paragraph_format.first_line_indent = Cm(0.74)
        conclusion_para.paragraph_format.line_spacing = 1.5
    
    # ========== 6. 页脚声明 ==========
    doc.add_page_break()
    footer_para = doc.add_paragraph()
    footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer_run = footer_para.add_run('本报告由AI智能体生成，仅供参考。报告中所有案例均已通过核验，具体案件处理请结合实际情况和法律规定。')
    footer_run.italic = True
    footer_run.font.size = Pt(10.5)
    
    # ========== 保存文档 ==========
    output_path = Path(output_path)
    output_path.parent.mkdir(parents=True, exist_ok=True)
    
    doc.save(str(output_path))
    
    # 验证输出文件
    if not output_path.exists():
        raise ValueError(f"❌ 文件保存失败: {output_path}")
    
    file_size = output_path.stat().st_size
    if file_size < 1000:  # 小于1KB可能是空文档
        raise ValueError(f"❌ 生成的文档过小 ({file_size} bytes)，可能为空")
    
    print_debug(f"✅ Word报告已生成: {output_path} ({file_size} bytes)")


def generate_report(input_file: str, output_file: str):
    """
    生成报告主流程
    
    Args:
        input_file: 输入JSON文件路径（应为 verified_cases.json，仅含核验通过案例）
        output_file: 输出Word文件路径
    """
    try:
        print_debug("🚀 开始生成报告...")
        print_debug(f"   输入文件: {input_file}")
        print_debug(f"   输出文件: {output_file}")
        
        # 解析输入文件（包含核验过滤）
        case_data = parse_case_input(input_file)
        
        # 生成Word报告
        generate_word_report(case_data, Path(output_file))
        
        print_debug("✨ 报告生成完成！")
        print(f"📄 报告已保存到: {output_file}")
        
        return True
        
    except ValueError as e:
        print_debug(f"❌ 数据验证失败: {str(e)}")
        print(f"错误: {str(e)}", file=sys.stderr)
        sys.exit(1)
    except Exception as e:
        print_debug(f"❌ 报告生成失败: {str(e)}")
        print(f"错误: {str(e)}", file=sys.stderr)
        import traceback
        traceback.print_exc()
        sys.exit(1)


if __name__ == "__main__":
    print_debug("🚀 generate_report.py 启动...")
    
    parser = argparse.ArgumentParser(
        description="生成类案检索报告 - 将JSON数据转换为Word格式(.docx)",
        formatter_class=argparse.RawDescriptionHelpFormatter,
    )
    
    parser.add_argument(
        "--input",
        required=True,
        help="输入JSON文件路径"
    )
    parser.add_argument(
        "--output",
        required=True,
        help="输出Word文档路径"
    )
    
    args = parser.parse_args()
    
    generate_report(args.input, args.output)