#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
财税分析报告生成器
支持生成HTML和Word格式的财税专业分析报告
"""

import json
import argparse
import sys
import os
import re
from pathlib import Path
from datetime import datetime
from typing import List, Optional, Tuple

try:
    from docx import Document
    from docx.shared import Pt, Inches, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
except ImportError:
    print("错误: 缺少依赖包，请安装 python-docx")
    print("安装命令: pip install python-docx==0.8.11")
    sys.exit(1)


def _split_markdown_blocks(text: str) -> List[dict]:
    if not text:
        return []

    lines = str(text).splitlines()
    blocks: List[dict] = []
    buf: List[str] = []
    i = 0

    def flush_buf():
        nonlocal buf
        if buf:
            blocks.append({'type': 'text', 'text': '\n'.join(buf)})
            buf = []

    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        if stripped.startswith('```'):
            flush_buf()
            language = stripped[3:].strip() or None
            code_lines: List[str] = []
            i += 1
            while i < len(lines) and not lines[i].strip().startswith('```'):
                code_lines.append(lines[i])
                i += 1
            blocks.append({'type': 'code', 'text': '\n'.join(code_lines), 'language': language})
            while i < len(lines) and not lines[i].strip().startswith('```'):
                i += 1
            i += 1
            continue

        if stripped == '':
            flush_buf()
            i += 1
            continue

        buf.append(line)
        i += 1

    flush_buf()
    return blocks


def _try_parse_list_block(text_block: str) -> Optional[Tuple[str, List[str]]]:
    lines = [line.strip() for line in str(text_block).splitlines() if line.strip()]
    if not lines:
        return None

    ul_items: List[str] = []
    for line in lines:
        m = re.match(r'^[-*]\s+(.*)$', line)
        if not m:
            ul_items = []
            break
        ul_items.append(m.group(1))
    if ul_items:
        return 'ul', ul_items

    ol_items: List[str] = []
    for line in lines:
        m = re.match(r'^\d+[.)]\s+(.*)$', line)
        if not m:
            ol_items = []
            break
        ol_items.append(m.group(1))
    if ol_items:
        return 'ol', ol_items

    return None


class HTMLReportGenerator:
    """HTML报告生成器"""

    def __init__(self, title, content):
        self.title = title
        self.content = content
        self._template_path = Path(__file__).resolve().parent.parent / 'assets' / 'html-template.html'

    def generate(self, output_path=None, return_content=False):
        """
        生成HTML报告

        Args:
            output_path: 输出文件路径，如果提供则保存到文件
            return_content: 是否返回HTML内容

        Returns:
            如果return_content=True，返回HTML字符串
            如果output_path提供，返回文件路径
        """
        html_content = self._build_html()

        # 如果需要返回内容
        if return_content:
            return html_content

        # 如果需要保存到文件
        if output_path:
            output_path = Path(output_path)
            output_path.parent.mkdir(parents=True, exist_ok=True)
            with open(output_path, 'w', encoding='utf-8') as f:
                f.write(html_content)
            return str(output_path)

    def _build_html(self):
        """构建HTML内容"""
        title_escaped = self._escape_html(self.title)
        body_parts: List[str] = []

        if 'summary' in self.content:
            body_parts.append('    <h2>摘要</h2>')
            body_parts.append(self._markdown_to_html(self.content['summary'], paragraph_style='text-indent: 2em;'))

        if 'sections' in self.content and self.content['sections']:
            for section in self.content['sections']:
                body_parts.append(self._render_section(section))

        if 'conclusion' in self.content:
            body_parts.append('    <h2>结论</h2>')
            body_parts.append(self._markdown_to_html(self.content['conclusion'], paragraph_style='text-indent: 2em;'))

        if 'references' in self.content and self.content['references']:
            body_parts.append(self._render_references())

        body_html = '\n'.join([part for part in body_parts if part])
        template_html = self._load_template_html()
        return template_html.replace('{{TITLE}}', title_escaped).replace('{{BODY}}', body_html)

    def _load_template_html(self) -> str:
        try:
            return self._template_path.read_text(encoding='utf-8')
        except OSError:
            title_escaped = self._escape_html(self.title)
            return f"""<!DOCTYPE html>
<html lang="zh-CN">
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>{title_escaped}</title>
</head>
<body>
    <h1>{title_escaped}</h1>
{{BODY}}
</body>
</html>"""

    def _render_section(self, section_data):
        """渲染章节"""
        title = section_data.get('title', '')
        html = f'    <h2>{self._escape_html(title)}</h2>\n'

        # 渲染表格
        if 'tables' in section_data:
            for table_data in section_data['tables']:
                html += self._render_table(table_data)

        # 渲染内容
        if 'content' in section_data:
            html += self._markdown_to_html(section_data['content'])

        # 渲染小节
        if 'subsections' in section_data:
            for subsection in section_data['subsections']:
                html += self._render_subsection(subsection)

        return html

    def _render_subsection(self, subsection):
        """渲染小节"""
        subtitle = subsection.get('subtitle', '')
        border_bottom = subsection.get('border_bottom', False)

        if border_bottom:
            html = f'    <h3 style="border-bottom: 1px dashed #ccc; padding-bottom: 0.5em;">{self._escape_html(subtitle)}</h3>\n'
        else:
            html = f'    <h3>{self._escape_html(subtitle)}</h3>\n'

        content = subsection.get('content', '')
        if content:
            html += self._markdown_to_html(content)

        # 渲染表格
        if 'tables' in subsection:
            for table_data in subsection['tables']:
                html += self._render_table(table_data)

        return html

    def _render_table(self, table_data):
        """渲染表格"""
        headers = table_data.get('headers', [])
        rows = table_data.get('rows', [])
        caption = table_data.get('caption', '')

        html = ''
        if caption:
            html += f'    <p style="text-align: center; font-weight: bold;">{self._escape_html(caption)}</p>\n'

        html += '    <table>\n'
        html += '        <thead>\n        <tr>\n'
        for header in headers:
            html += f'            <th>{self._escape_html(str(header))}</th>\n'
        html += '        </tr>\n        </thead>\n'
        html += '        <tbody>\n'

        for row in rows:
            html += '        <tr>\n'
            for cell in row:
                html += f'            <td>{self._escape_html(str(cell))}</td>\n'
            html += '        </tr>\n'

        html += '        </tbody>\n'
        html += '    </table>\n'

        return html

    def _render_references(self):
        """渲染参考文献"""
        html = '    <h2>参考文献</h2>\n'
        html += '    <div class="references">\n'
        html += '        <ol>\n'

        for ref in self.content['references']:
            title = ref.get('title', '')
            url = ref.get('url', '')
            if url:
                html += f'            <li><a href="{self._escape_html(url)}" target="_blank" rel="noopener noreferrer">{self._escape_html(title)}</a></li>\n'
            else:
                html += f'            <li>{self._escape_html(title)}</li>\n'

        html += '        </ol>\n'
        html += '    </div>\n'

        return html

    def _escape_html(self, text):
        """转义HTML特殊字符"""
        if not text:
            return ''
        text = str(text)
        text = text.replace('&', '&amp;')
        text = text.replace('<', '&lt;')
        text = text.replace('>', '&gt;')
        text = text.replace('"', '&quot;')
        text = text.replace("'", '&#39;')
        return text

    def _convert_inline_markdown_to_html(self, text: str) -> str:
        escaped = self._escape_html(text)
        result: List[str] = []
        i = 0
        bold = False
        italic = False
        code = False

        while i < len(escaped):
            if not code and escaped[i:i + 2] == '**':
                result.append('</strong>' if bold else '<strong>')
                bold = not bold
                i += 2
                continue
            if not code and escaped[i] == '*':
                result.append('</em>' if italic else '<em>')
                italic = not italic
                i += 1
                continue
            if escaped[i] == '`':
                result.append('</code>' if code else '<code>')
                code = not code
                i += 1
                continue
            result.append(escaped[i])
            i += 1

        if code:
            result.append('</code>')
        if italic:
            result.append('</em>')
        if bold:
            result.append('</strong>')

        return ''.join(result)

    def _markdown_to_html(self, text, paragraph_style: Optional[str] = None) -> str:
        blocks = _split_markdown_blocks(str(text) if text is not None else '')
        html_parts: List[str] = []

        for block in blocks:
            if block['type'] == 'code':
                code_text = self._escape_html(block.get('text', ''))
                html_parts.append(f'    <pre><code>{code_text}</code></pre>')
                continue

            block_text = block.get('text', '')
            list_result = _try_parse_list_block(block_text)
            if list_result:
                list_type, items = list_result
                tag = 'ul' if list_type == 'ul' else 'ol'
                html_parts.append(f'    <{tag}>')
                for item in items:
                    html_parts.append(f'        <li>{self._convert_inline_markdown_to_html(item)}</li>')
                html_parts.append(f'    </{tag}>')
                continue

            lines = [line.rstrip() for line in str(block_text).splitlines()]
            inner = '<br>'.join(self._convert_inline_markdown_to_html(line) for line in lines if line.strip() != '')
            style_attr = f' style="{paragraph_style}"' if paragraph_style else ''
            html_parts.append(f'    <p{style_attr}>{inner}</p>')

        return '\n'.join(html_parts) + '\n'


class WordReportGenerator:
    """Word报告生成器"""

    def __init__(self, title, content):
        self.title = title
        self.content = content
        self.doc = Document()

    def generate(self, output_path, return_file_path=True):
        """
        生成Word报告

        Args:
            output_path: 输出文件路径
            return_file_path: 是否返回文件路径

        Returns:
            如果return_file_path=True，返回文件路径
        """
        self._setup_document()
        self._add_title()
        self._add_summary()
        self._add_sections()
        self._add_conclusion()
        self._add_references()

        # 保存文档
        output_path = Path(output_path)
        output_path.parent.mkdir(parents=True, exist_ok=True)
        self.doc.save(str(output_path))

        if return_file_path:
            return str(output_path)

    def _setup_document(self):
        """设置文档基本格式"""
        # 设置页面边距
        sections = self.doc.sections
        for section in sections:
            section.top_margin = Inches(1.0)
            section.bottom_margin = Inches(1.0)
            section.left_margin = Inches(1.25)
            section.right_margin = Inches(1.25)

    def _add_title(self):
        """添加标题"""
        title = self.doc.add_paragraph(self.title)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_run = title.runs[0]
        title_run.font.name = '黑体'
        title_run.font.size = Pt(16)
        title_run.font.bold = True
        title_run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')

    def _add_summary(self):
        """添加摘要"""
        if 'summary' in self.content:
            self.doc.add_paragraph('摘要', style='Heading 2')
            self._add_markdown_blocks(self.content['summary'])

    def _add_sections(self):
        """添加各个章节（灵活结构）"""
        if 'sections' in self.content and self.content['sections']:
            for section in self.content['sections']:
                self._render_section(section)

    def _render_section(self, section_data):
        """渲染章节"""
        # 添加H2标题
        title = section_data.get('title', '')
        h2 = self.doc.add_paragraph(title, style='Heading 2')
        h2_run = h2.runs[0]
        h2_run.font.name = '黑体'
        h2_run.font.size = Pt(15)
        h2_run.font.bold = True
        h2_run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')

        # 渲染内容
        if 'content' in section_data:
            self._add_markdown_blocks(section_data['content'])

        # 渲染表格
        if 'tables' in section_data:
            for table_data in section_data['tables']:
                self._render_table(table_data)

        # 渲染小节
        if 'subsections' in section_data:
            for subsection in section_data['subsections']:
                self._render_subsection(subsection)

    def _render_subsection(self, subsection):
        """渲染小节"""
        subtitle = subsection.get('subtitle', '')

        # 添加H3标题
        h3 = self.doc.add_paragraph(subtitle, style='Heading 3')
        h3_run = h3.runs[0]
        h3_run.font.name = '黑体'
        h3_run.font.size = Pt(14)
        h3_run.font.bold = True
        h3_run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')

        # 设置段落格式
        h3_format = h3.paragraph_format
        h3_format.space_before = Pt(6)
        h3_format.space_after = Pt(6)

        # 添加内容
        content = subsection.get('content', '')
        if content:
            self._add_markdown_blocks(content)

        # 渲染表格
        if 'tables' in subsection:
            for table_data in subsection['tables']:
                self._render_table(table_data)

    def _render_table(self, table_data):
        """渲染表格"""
        headers = table_data.get('headers', [])
        rows = table_data.get('rows', [])
        caption = table_data.get('caption', '')

        if caption:
            caption_para = self.doc.add_paragraph(caption)
            caption_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            caption_run = caption_para.runs[0]
            caption_run.font.bold = True

        # 创建表格
        table = self.doc.add_table(rows=len(rows) + 1, cols=len(headers))
        table.style = 'Table Grid'

        # 设置表头
        header_cells = table.rows[0].cells
        for i, header in enumerate(headers):
            cell = header_cells[i]
            cell.text = str(header)
            for run in cell.paragraphs[0].runs:
                run.font.bold = True
                run.font.size = Pt(10.5)
                run.font.name = '宋体'
                run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

        # 设置数据行
        for i, row in enumerate(rows):
            row_cells = table.rows[i + 1].cells
            for j, cell_data in enumerate(row):
                cell = row_cells[j]
                cell.text = str(cell_data)
                for run in cell.paragraphs[0].runs:
                    run.font.size = Pt(10.5)
                    run.font.name = '宋体'
                    run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

        # 设置表格样式
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
                    paragraph.space_before = Pt(2)
                    paragraph.space_after = Pt(2)

    def _add_conclusion(self):
        """添加结论"""
        if 'conclusion' in self.content:
            self.doc.add_paragraph('结论', style='Heading 2')
            self._add_markdown_blocks(self.content['conclusion'])

    def _add_markdown_blocks(self, text):
        blocks = _split_markdown_blocks(str(text) if text is not None else '')
        for block in blocks:
            if block['type'] == 'code':
                para = self.doc.add_paragraph()
                run = para.add_run(block.get('text', ''))
                run.font.name = 'Consolas'
                run.font.size = Pt(10.5)
                self._set_paragraph_format(para)
                continue

            block_text = block.get('text', '')
            list_result = _try_parse_list_block(block_text)
            if list_result:
                list_type, items = list_result
                style = 'List Bullet' if list_type == 'ul' else 'List Number'
                for item in items:
                    para = self.doc.add_paragraph(style=style)
                    self._add_inline_markdown_runs(para, item)
                    self._set_paragraph_format(para)
                continue

            para = self.doc.add_paragraph()
            self._add_inline_markdown_runs(para, block_text)
            self._set_paragraph_format(para)

    def _add_inline_markdown_runs(self, para, text: str):
        raw = str(text or '')
        i = 0
        bold = False
        italic = False
        code = False
        buf: List[str] = []

        def flush():
            nonlocal buf
            if not buf:
                return
            run = para.add_run(''.join(buf))
            run.bold = bold if not code else False
            run.italic = italic if not code else False
            if code:
                run.font.name = 'Consolas'
                run.font.size = Pt(10.5)
            buf = []

        while i < len(raw):
            if not code and raw[i:i + 2] == '**':
                flush()
                bold = not bold
                i += 2
                continue
            if not code and raw[i] == '*':
                flush()
                italic = not italic
                i += 1
                continue
            if raw[i] == '`':
                flush()
                code = not code
                i += 1
                continue
            buf.append(raw[i])
            i += 1

        flush()

    def _add_references(self):
        """添加参考文献"""
        if 'references' in self.content and self.content['references']:
            self.doc.add_paragraph('参考文献', style='Heading 2')

            for i, ref in enumerate(self.content['references'], 1):
                title = ref.get('title', '')
                url = ref.get('url', '')

                ref_text = f'[{i}] {title}'
                if url:
                    ref_text += f'\n    {url}'

                para = self.doc.add_paragraph(ref_text)
                para.paragraph_format.left_indent = Inches(0.25)
                para.paragraph_format.hanging_indent = Inches(-0.25)
                self._set_paragraph_format(para)

    def _set_paragraph_format(self, para):
        """设置段落格式"""
        para.paragraph_format.line_spacing = 1.5
        para.paragraph_format.space_before = Pt(6)
        para.paragraph_format.space_after = Pt(6)
        para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        style_name = getattr(getattr(para, 'style', None), 'name', '') or ''
        if style_name.startswith('List'):
            para.paragraph_format.first_line_indent = Inches(0)
        else:
            para.paragraph_format.first_line_indent = Inches(0.3)  # 2字符缩进

        # 设置字体
        for run in para.runs:
            run.font.size = Pt(12)
            run.font.name = '宋体'
            run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

    def _split_paragraphs(self, text):
        """分割段落"""
        if not text:
            return []
        # 按两个或更多换行符分割段落
        paragraphs = text.split('\n\n')
        return [p.strip() for p in paragraphs if p.strip()]


def main():
    """主函数"""
    parser = argparse.ArgumentParser(description='财税分析报告生成器')
    parser.add_argument('--format', required=True, choices=['html', 'word', 'both'], help='输出格式：html、word 或 both')
    parser.add_argument('--title', required=True, help='报告标题')
    parser.add_argument('--output', help='输出文件路径（不含扩展名）')
    parser.add_argument('--content', required=True, help='报告内容（JSON格式字符串）')

    args = parser.parse_args()

    try:
        content = json.loads(args.content)
    except json.JSONDecodeError as e:
        print(f"错误: JSON格式不正确 - {e}")
        sys.exit(1)

    results = {}

    # 生成HTML
    if args.format in ['html', 'both']:
        if args.output:
            html_path = f"{args.output}.html"
        else:
            html_path = f"./report_{datetime.now().strftime('%Y%m%d_%H%M%S')}.html"

        html_generator = HTMLReportGenerator(args.title, content)
        html_result = html_generator.generate(html_path)
        results['html'] = html_result
        print(f"HTML报告已生成: {html_result}")

    # 生成Word
    if args.format in ['word', 'both']:
        if args.output:
            word_path = f"{args.output}.docx"
        else:
            word_path = f"./report_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"

        word_generator = WordReportGenerator(args.title, content)
        word_result = word_generator.generate(word_path)
        results['word'] = word_result
        print(f"Word报告已生成: {word_result}")

    # 输出结果（用于脚本调用）
    if results:
        print(f"RESULT: {json.dumps(results, ensure_ascii=False)}")


if __name__ == '__main__':
    main()
